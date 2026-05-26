#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件解析模块 - 支持 Word (.docx) 和 TXT 文件
"""
import re
from pathlib import Path
from typing import Optional, Dict, Any
from abc import ABC, abstractmethod


class DocumentParser(ABC):
    """文档解析基类"""

    @abstractmethod
    def parse(self, file_path: str) -> str:
        """解析文件并返回文本内容"""
        pass

    @staticmethod
    def clean_text(text: str) -> str:
        """清理文本"""
        # 移除多余空白
        text = re.sub(r'\s+', ' ', text)
        # 移除特殊控制字符
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        return text.strip()


class TxtParser(DocumentParser):
    """TXT文件解析器"""

    def parse(self, file_path: str) -> str:
        """解析TXT文件"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()

        return self.clean_text(content)


class WordParser(DocumentParser):
    """Word文件解析器"""

    def parse(self, file_path: str) -> str:
        """解析Word文件"""
        try:
            import docx
        except ImportError:
            raise ImportError(
                "请安装 python-docx: pip install python-docx\n"
                "或者将Word文件另存为TXT格式使用"
            )

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        doc = docx.Document(path)

        # 提取所有段落
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        content = '\n'.join(paragraphs)
        return self.clean_text(content)


class DocumentParserFactory:
    """文档解析工厂"""

    _parsers = {
        '.txt': TxtParser,
        '.docx': WordParser,
    }

    @classmethod
    def get_parser(cls, file_path: str) -> DocumentParser:
        """根据文件扩展名获取解析器"""
        ext = Path(file_path).suffix.lower()
        parser_class = cls._parsers.get(ext)
        if parser_class is None:
            supported = ', '.join(cls._parsers.keys())
            raise ValueError(f"不支持的文件格式: {ext}，支持的格式: {supported}")
        return parser_class()

    @classmethod
    def parse(cls, file_path: str) -> str:
        """直接解析文件"""
        parser = cls.get_parser(file_path)
        return parser.parse(file_path)


def parse_document(file_path: str) -> str:
    """
    解析文档文件（入口函数）

    支持:
    - .txt: 纯文本文件
    - .docx: Word文档

    Args:
        file_path: 文件路径

    Returns:
        解析后的文本内容
    """
    return DocumentParserFactory.parse(file_path)


def extract_topic_from_text(text: str) -> Dict[str, Any]:
    """
    从文本中提取话题信息

    Args:
        text: 文本内容

    Returns:
        包含话题和关键词的字典
    """
    # 简单的关键词提取（基于词频）
    # 实际应用中可用更复杂的NLP方法
    import jieba
    from collections import Counter

    # 停用词
    stopwords = {
        '的', '了', '是', '在', '我', '有', '和', '就', '不', '人', '都', '一', '一个',
        '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看', '好',
        '自己', '这', '那', '什么', '怎么', '为什么', '如何', '可以', '这个', '那个',
        '啊', '呢', '吧', '吗', '哦', '嗯', '呀', '哈', '呐', '么', '哪些', '以及',
        '对', '于', '与', '或', '但', '而', '所以', '因为', '如果', '虽然', '然而',
        '通过', '对于', '关于', '以及', '包括', '能够', '可以', '应该', '必须', '需要',
        '进行', '使用', '通过', '根据', '按照', '为了', '由于', '关于', '随着',
    }

    words = jieba.cut(text)
    # 过滤停用词和短词
    filtered = [w for w in words if w not in stopwords and len(w) >= 2]
    word_counts = Counter(filtered)

    # 取前10个高频词作为关键词
    keywords = [word for word, count in word_counts.most_common(10)]

    return {
        'topic': text[:200],  # 取前200字符作为话题描述
        'keywords': keywords,
        'full_text': text,
    }


if __name__ == "__main__":
    # 测试
    import sys
    if len(sys.argv) > 1:
        try:
            content = parse_document(sys.argv[1])
            print(f"解析成功，字符数: {len(content)}")
            print(f"前200字符: {content[:200]}...")
        except Exception as e:
            print(f"解析失败: {e}")
    else:
        print("用法: python document_parser.py <文件路径>")
