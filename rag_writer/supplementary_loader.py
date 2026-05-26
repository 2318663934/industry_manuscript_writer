#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
补充资料加载模块 - 支持从多种来源提取内容

支持格式:
- 链接: 微信文章链接等网页内容
- 文件: TXT, DOCX, PDF, Excel (.xlsx/.xls)
"""
import re
import time
import requests
from pathlib import Path
from typing import Optional, Dict, Any, List, Union
from dataclasses import dataclass, field
from bs4 import BeautifulSoup
from requests.exceptions import RequestException

from .config import settings


@dataclass
class SupplementaryMaterial:
    """补充资料"""
    title: str
    source: str  # 来源描述：文件名、URL、或"用户上传"
    source_type: str  # 'url', 'file', 'text'
    content: str
    author: str = ""
    date: str = ""

    @property
    def content_length(self) -> int:
        return len(self.content)

    def to_knowledge_context(self) -> Dict[str, str]:
        """转换为KnowledgeContext格式"""
        return {
            "title": self.title,
            "url": self.source if self.source_type == 'url' else '',
            "author": self.author,
            "date": self.date,
            "content": self.content,
            "source_type": "supplementary",  # 标记为补充资料
        }


class BaseParser:
    """解析器基类"""

    def parse(self, path: str) -> str:
        raise NotImplementedError

    @staticmethod
    def clean_text(text: str) -> str:
        """清理文本"""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        return text.strip()


class TxtParser(BaseParser):
    """TXT文件解析器"""

    def parse(self, path: str) -> str:
        path_obj = Path(path)
        if not path_obj.exists():
            raise FileNotFoundError(f"文件不存在: {path}")

        with open(path_obj, 'r', encoding='utf-8') as f:
            content = f.read()

        return self.clean_text(content)


class WordParser(BaseParser):
    """Word文件解析器 (.docx)"""

    def parse(self, path: str) -> str:
        try:
            import docx
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        path_obj = Path(path)
        if not path_obj.exists():
            raise FileNotFoundError(f"文件不存在: {path}")

        doc = docx.Document(path_obj)
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        content = '\n'.join(paragraphs)
        return self.clean_text(content)


class PdfParser(BaseParser):
    """PDF文件解析器"""

    def parse(self, path: str) -> str:
        try:
            import pdfplumber
        except ImportError:
            try:
                import PyPDF2
            except ImportError:
                raise ImportError(
                    "请安装 pdfplumber: pip install pdfplumber\n"
                    "或者 pip install PyPDF2"
                )

        path_obj = Path(path)
        if not path_obj.exists():
            raise FileNotFoundError(f"文件不存在: {path}")

        # 优先使用 pdfplumber（支持表格）
        try:
            import pdfplumber
            texts = []
            with pdfplumber.open(path_obj) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        texts.append(page_text)
                    # 尝试提取表格
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            row_text = ' | '.join(str(cell) if cell else '' for cell in row)
                            if row_text.strip():
                                texts.append(row_text)
            return self.clean_text('\n'.join(texts))
        except ImportError:
            pass

        # Fallback to PyPDF2
        import PyPDF2
        texts = []
        with open(path_obj, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    texts.append(page_text)
        return self.clean_text('\n'.join(texts))


class ExcelParser(BaseParser):
    """Excel文件解析器 (.xlsx/.xls)"""

    def parse(self, path: str) -> str:
        try:
            import openpyxl
        except ImportError:
            raise ImportError("请安装 openpyxl: pip install openpyxl")

        path_obj = Path(path)
        if not path_obj.exists():
            raise FileNotFoundError(f"文件不存在: {path}")

        texts = []
        wb = openpyxl.load_workbook(path_obj, data_only=True)

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            texts.append(f"=== 工作表: {sheet_name} ===")

            for row in sheet.iter_rows(values_only=True):
                row_text = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                if row_text.strip():
                    texts.append(row_text)

            texts.append("")

        wb.close()
        return self.clean_text('\n'.join(texts))


class URLParser:
    """URL链接解析器（支持微信文章等网页）"""

    def __init__(self, timeout: int = 30):
        self.timeout = timeout
        self.session = requests.Session()
        self.session.headers.update({
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        })

    def parse(self, url: str) -> SupplementaryMaterial:
        """解析URL并提取内容"""
        try:
            response = self.session.get(url, timeout=self.timeout)
            response.raise_for_status()

            # 根据URL类型选择解析方法
            if 'mp.weixin.qq.com' in url:
                return self._parse_wechat_article(url, response.text)
            else:
                return self._parse_generic_page(url, response.text)

        except RequestException as e:
            raise RuntimeError(f"获取链接内容失败: {e}")

    def _parse_wechat_article(self, url: str, html: str) -> SupplementaryMaterial:
        """解析微信文章"""
        soup = BeautifulSoup(html, "lxml")

        title = ""
        title_elem = soup.select_one("h1#activity-name")
        if not title_elem:
            title_elem = soup.select_one("h1.rich_media_title")
        if title_elem:
            title = title_elem.get_text(strip=True)

        author = ""
        author_elem = soup.select_one("span#js_name")
        if not author_elem:
            author_elem = soup.select_one("a.rich_media_meta.rich_media_meta_link")
        if author_elem:
            author = author_elem.get_text(strip=True)

        date = ""
        date_elem = soup.select_one("span#publish_time")
        if not date_elem:
            date_elem = soup.select_one("em#post-date")
        if date_elem:
            date = date_elem.get_text(strip=True)

        content_elem = soup.select_one("div#js_content")
        content_text = ""
        if content_elem:
            content_text = content_elem.get_text(separator="\n", strip=True)

        return SupplementaryMaterial(
            title=title or "微信文章",
            source=url,
            source_type="url",
            content=self.clean_text(content_text),
            author=author,
            date=date,
        )

    def _parse_generic_page(self, url: str, html: str) -> SupplementaryMaterial:
        """解析通用网页"""
        soup = BeautifulSoup(html, "lxml")

        # 移除脚本和样式
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        # 提取标题
        title = ""
        title_elem = soup.find("title")
        if title_elem:
            title = title_elem.get_text(strip=True)
        if not title:
            h1 = soup.find("h1")
            if h1:
                title = h1.get_text(strip=True)

        # 提取正文（尝试找主要内容区域）
        content_elem = soup.find("main") or soup.find("article") or soup.find("div", class_=re.compile(r"content|article|post|entry"))
        if not content_elem:
            content_elem = soup.find("body")

        content_text = ""
        if content_elem:
            content_text = content_elem.get_text(separator="\n", strip=True)

        return SupplementaryMaterial(
            title=title or "网页文章",
            source=url,
            source_type="url",
            content=self.clean_text(content_text[:10000] if len(content_text) > 10000 else content_text),  # 限制长度
            author="",
            date="",
        )


class TextInputParser:
    """纯文本输入解析器（用户直接粘贴的文本）"""

    def parse(self, text: str, title: str = "用户输入文本") -> SupplementaryMaterial:
        return SupplementaryMaterial(
            title=title,
            source="用户输入",
            source_type="text",
            content=self.clean_text(text),
            author="",
            date="",
        )


class SupplementaryLoader:
    """
    补充资料加载器

    支持加载:
    - 文件路径: TXT, DOCX, PDF, Excel
    - URL链接: 微信文章等网页
    - 纯文本: 用户直接粘贴的文本
    """

    # 文件扩展名到解析器的映射
    FILE_PARSERS = {
        '.txt': TxtParser,
        '.docx': WordParser,
        '.pdf': PdfParser,
        '.xlsx': ExcelParser,
        '.xls': ExcelParser,
    }

    def __init__(self):
        self.url_parser = URLParser()
        self.text_parser = TextInputParser()

    def load(self, source: Union[str, List[str]]) -> List[SupplementaryMaterial]:
        """
        加载补充资料

        Args:
            source: 可以是:
                - 单个文件路径/URL/文本
                - 列表: [文件路径1, URL1, "纯文本", ...]

        Returns:
            SupplementaryMaterial列表
        """
        if isinstance(source, str):
            source = [source]

        materials = []
        for item in source:
            try:
                material = self._load_single(item)
                if material and material.content:
                    materials.append(material)
            except Exception as e:
                print(f"警告: 加载补充资料失败 '{item}': {e}")
                continue

        return materials

    def _load_single(self, source: str) -> Optional[SupplementaryMaterial]:
        """加载单个补充资料"""
        source = source.strip()
        if not source:
            return None

        # 判断类型
        if self._is_url(source):
            return self.url_parser.parse(source)
        elif self._is_file_path(source):
            return self._load_file(source)
        else:
            # 假定是纯文本
            return self.text_parser.parse(source)

    def _is_url(self, source: str) -> bool:
        """判断是否是URL"""
        return source.startswith(('http://', 'https://'))

    def _is_file_path(self, source: str) -> bool:
        """判断是否是文件路径"""
        path = Path(source)
        # 检查文件是否存在
        if path.exists():
            return True
        # 检查扩展名
        if path.suffix.lower() in self.FILE_PARSERS:
            return True
        return False

    def _load_file(self, file_path: str) -> SupplementaryMaterial:
        """加载文件"""
        path = Path(file_path)
        suffix = path.suffix.lower()

        parser_class = self.FILE_PARSERS.get(suffix)
        if parser_class is None:
            raise ValueError(f"不支持的文件格式: {suffix}")

        parser = parser_class()
        content = parser.parse(str(path))

        return SupplementaryMaterial(
            title=path.stem,  # 文件名（不含扩展名）
            source=str(path.absolute()),
            source_type="file",
            content=content,
            author="",
            date="",
        )


def load_supplementary_materials(
    sources: Union[str, List[str]]
) -> List[SupplementaryMaterial]:
    """
    便捷函数：加载补充资料

    Args:
        sources: 文件路径/URL/文本，或它们的列表

    Returns:
        SupplementaryMaterial列表
    """
    loader = SupplementaryLoader()
    return loader.load(sources)


if __name__ == "__main__":
    # 测试
    import sys

    if len(sys.argv) > 1:
        sources = sys.argv[1:]
        loader = SupplementaryLoader()
        materials = loader.load(sources)

        print(f"成功加载 {len(materials)} 份补充资料:")
        for i, m in enumerate(materials, 1):
            print(f"\n--- 资料 {i} ---")
            print(f"标题: {m.title}")
            print(f"来源: {m.source} ({m.source_type})")
            print(f"字数: {m.content_length}")
            print(f"内容预览: {m.content[:200]}...")
    else:
        print("用法: python supplementary_loader.py <文件路径/URL/文本> ...")
